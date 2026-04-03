"""
File processing engine for Lazy Footers.

This module contains ALL file manipulation logic:
  - PDF footer replacement   (via PyMuPDF / fitz — redacts old + stamps 3-column)
  - DOCX footer editing      (via python-docx — 3-column table footer)
  - DOCX → PDF conversion    (via docx2pdf — uses MS Word on Windows)
  - Single-file pipeline orchestration

Output is ALWAYS PDF regardless of input format.
Footer layout: [left text]    [center text]    [right text]
"""

import logging
import os
import shutil
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

FOOTER_ZONE_HEIGHT = 90  # pt wiped from bottom of each PDF page


# ---------------------------------------------------------------------------
# PDF footer replacement (via PyMuPDF) — 3-column layout
# ---------------------------------------------------------------------------

def add_footer_to_pdf(
    pdf_path: str,
    output_path: str,
    left_text: str,
    center_text: str,
    right_text: str,
) -> None:
    """
    Replace the footer zone on every page with a 3-column footer.

    Layout:  [left_text]        [center_text]        [right_text]

    Steps:
      1. Erase the bottom zone of each page (redaction)
      2. Stamp left-aligned text on the left third
      3. Stamp center-aligned text in the middle third
      4. Stamp right-aligned text on the right third
    """
    try:
        document = fitz.open(pdf_path)
    except Exception as error:
        raise RuntimeError(
            f"Failed to open PDF: {pdf_path}. Error: {error}"
        ) from error

    try:
        for page in document:
            pr = page.rect

            # --- Step 1: Redact the footer zone ---
            footer_zone = fitz.Rect(0, pr.height - FOOTER_ZONE_HEIGHT, pr.width, pr.height)
            page.add_redact_annot(footer_zone, fill=(1, 1, 1))
            page.apply_redactions()

            # Shared vertical positioning within the cleared zone
            # Stamp vertically centered in standard footer space (bottom 0.5 to 1 inch)
            top = pr.height - 45
            bottom = pr.height - 15
            margin = 30

            # --- Step 2: Left text (left-aligned) ---
            if left_text.strip():
                box_left = fitz.Rect(margin, top, pr.width / 3, bottom)
                page.insert_textbox(
                    box_left, left_text,
                    fontsize=9, fontname="helv",
                    color=(0.25, 0.25, 0.25),
                    align=fitz.TEXT_ALIGN_LEFT,
                )

            # --- Step 3: Center text (center-aligned) ---
            if center_text.strip():
                box_center = fitz.Rect(pr.width / 3, top, 2 * pr.width / 3, bottom)
                page.insert_textbox(
                    box_center, center_text,
                    fontsize=9, fontname="helv",
                    color=(0.25, 0.25, 0.25),
                    align=fitz.TEXT_ALIGN_CENTER,
                )

            # --- Step 4: Right text (right-aligned) ---
            if right_text.strip():
                box_right = fitz.Rect(2 * pr.width / 3, top, pr.width - margin, bottom)
                page.insert_textbox(
                    box_right, right_text,
                    fontsize=9, fontname="helv",
                    color=(0.25, 0.25, 0.25),
                    align=fitz.TEXT_ALIGN_RIGHT,
                )

        document.save(output_path)
        logger.info("3-column footer set on %d pages: %s", len(document), output_path)

    except Exception as error:
        raise RuntimeError(f"Failed to stamp footer: {error}") from error
    finally:
        document.close()


# ---------------------------------------------------------------------------
# DOCX footer editing (via python-docx) — 3-column table
# ---------------------------------------------------------------------------

def set_footer_in_docx(
    docx_path: str,
    left_text: str,
    center_text: str,
    right_text: str,
) -> None:
    """
    Set a 3-part footer on every section using Word's tab-stop technique.

    One paragraph, two tab stops (center + right), text separated by tabs:
        [left_text] TAB [center_text] TAB [right_text]

    Tab positions are derived from the section's actual printable width
    (page_width - left_margin - right_margin) so no clipping occurs on
    any page size or margin setting.
    """
    # 1 twip = 635 EMU (Word's internal unit)
    EMU_PER_TWIP = 635

    try:
        document = Document(docx_path)
    except Exception as error:
        raise RuntimeError(f"Failed to open DOCX: {docx_path}. Error: {error}") from error

    for section in document.sections:
        # --- Calculate actual printable text width in twips ---
        page_w   = section.page_width   or (12240 * EMU_PER_TWIP)
        margin_l = section.left_margin  or (1440  * EMU_PER_TWIP)
        margin_r = section.right_margin or (1440  * EMU_PER_TWIP)
        text_width_twips = int((page_w - margin_l - margin_r) / EMU_PER_TWIP)

        # Word has 3 different footer types per section. Wipe and replace ALL of them.
        for footer_attr in ['footer', 'first_page_footer', 'even_page_footer']:
            footer = getattr(section, footer_attr, None)
            if footer is None:
                continue

            footer.is_linked_to_previous = False

            # --- Clear all existing footer content aggressively ---
            for para in footer.paragraphs:
                para.clear()
            for child in list(footer._element):
                footer._element.remove(child)

            # --- Build new paragraph with 2 tab stops ---
            para = footer.add_paragraph()

            # Set tab stops via XML
            pPr = para._p.get_or_add_pPr()
            tabs_el = OxmlElement("w:tabs")

            ctr = OxmlElement("w:tab")
            ctr.set(qn("w:val"), "center")
            ctr.set(qn("w:pos"), str(text_width_twips // 2))
            tabs_el.append(ctr)

            rgt = OxmlElement("w:tab")
            rgt.set(qn("w:val"), "right")
            rgt.set(qn("w:pos"), str(text_width_twips))
            tabs_el.append(rgt)

            pPr.append(tabs_el)

            # --- Fill in the three columns ---
            def _add_run(text: str) -> None:
                r = para.add_run(text)
                r.font.size = Pt(9)

            _add_run(left_text or "")
            _add_run("\t")
            _add_run(center_text or "")
            _add_run("\t")
            _add_run(right_text or "")

    document.save(docx_path)
    logger.info("DOCX footer set (%d sections): %s", len(document.sections), docx_path)


# ---------------------------------------------------------------------------
# Single-file pipeline — no format conversion
# ---------------------------------------------------------------------------

def process_single_file(
    saved_path: str,
    original_name: str,
    footer_left: str,
    footer_center: str,
    footer_right: str,
    output_directory: str,
    temp_directory: str,
) -> str:
    """
    Process one file and return the output path.
      - PDF  → wipe old footer, stamp 3-column → output as PDF
      - DOCX → set 3-part footer via tab stops  → output as DOCX
    No format conversion. Both operations are instant.
    """
    ext  = Path(original_name).suffix.lower()
    stem = Path(original_name).stem

    if ext == ".pdf":
        output_path = os.path.join(output_directory, f"{stem}_with_footer.pdf")
        add_footer_to_pdf(saved_path, output_path, footer_left, footer_center, footer_right)

    elif ext == ".docx":
        output_path = os.path.join(output_directory, f"{stem}_with_footer.docx")
        shutil.copy2(saved_path, output_path)
        set_footer_in_docx(output_path, footer_left, footer_center, footer_right)

    else:
        raise ValueError(f"Unsupported type '{ext}'. Only .pdf and .docx accepted.")

    logger.info("Done: '%s' → '%s'", original_name, output_path)
    return output_path
